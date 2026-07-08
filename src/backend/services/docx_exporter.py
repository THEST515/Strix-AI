from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from src.backend.domain.models import Finding, ScanReport, ScanTask


SEVERITY_LABELS = {
    "critical": "严重",
    "high": "高危",
    "medium": "中危",
    "low": "低危",
    "info": "提示",
}


def build_docx_export_bytes(
    *,
    template_path: str | Path,
    task_exports: list[dict[str, object]],
) -> bytes:
    template_document = Document(str(template_path))
    target_document = Document(str(template_path))

    metadata = _collect_template_metadata(template_document)
    _fill_section(
        target_document,
        metadata,
        paragraph_offset=0,
        table_offset=0,
        task_export=task_exports[0],
    )

    for task_export in task_exports[1:]:
        target_document.add_page_break()
        paragraph_offset = len(target_document.paragraphs)
        table_offset = len(target_document.tables)
        _append_template_section(target_document, template_document)
        _fill_section(
            target_document,
            metadata,
            paragraph_offset=paragraph_offset,
            table_offset=table_offset,
            task_export=task_export,
        )

    output = BytesIO()
    target_document.save(output)
    return output.getvalue()


def _append_template_section(target_document: DocumentObject, template_document: DocumentObject) -> int:
    body = target_document.element.body
    sect_pr = body.sectPr

    for child in template_document.element.body.iterchildren():
        if child.tag.endswith("sectPr"):
            continue
        body.insert(body.index(sect_pr), deepcopy(child))


def _collect_template_metadata(template_document: DocumentObject) -> dict[str, object]:
    system_name_index = None
    risk_heading_indices: list[int] = []

    for index, paragraph in enumerate(template_document.paragraphs):
        text = paragraph.text.strip()
        if text == "XX系统":
            system_name_index = index
        elif text == "[X风险]风险名称":
            risk_heading_indices.append(index)

    if system_name_index is None or not risk_heading_indices:
        raise ValueError("DOCX 模板缺少系统名或风险标题占位段落")

    return {
        "paragraph_count": len(template_document.paragraphs),
        "table_count": len(template_document.tables),
        "system_name_index": system_name_index,
        "risk_heading_indices": risk_heading_indices,
    }

def _fill_section(
    document: DocumentObject,
    metadata: dict[str, object],
    *,
    paragraph_offset: int,
    table_offset: int,
    task_export: dict[str, object],
) -> None:
    system_name_index = int(metadata["system_name_index"])
    base_risk_heading_indices = list(metadata["risk_heading_indices"])

    task = task_export["task"]
    report = task_export["report"]
    summary = task_export["summary"]
    if not isinstance(task, ScanTask) or not isinstance(report, ScanReport) or not isinstance(summary, dict):
        raise TypeError("task_exports entries must contain ScanTask, ScanReport, and summary dict")

    _replace_paragraph_text(document.paragraphs[paragraph_offset + system_name_index], task.name)

    findings = report.findings if report.findings else [_build_no_findings_placeholder(task, summary)]
    risk_heading_indices = _ensure_capacity_for_findings(
        document,
        findings=findings,
        paragraph_offset=paragraph_offset,
        table_offset=table_offset,
        base_risk_heading_indices=base_risk_heading_indices,
    )

    for finding_index, finding in enumerate(findings):
        heading_paragraph = document.paragraphs[paragraph_offset + risk_heading_indices[finding_index]]
        _replace_paragraph_text(
            heading_paragraph,
            f"[{SEVERITY_LABELS.get(finding.severity, finding.severity)}]{finding.title}",
        )
        _fill_risk_table(
            document.tables[table_offset + finding_index],
            task=task,
            finding=finding,
            summary=summary,
        )

    for remaining_index in range(len(findings), len(risk_heading_indices)):
        _replace_paragraph_text(document.paragraphs[paragraph_offset + risk_heading_indices[remaining_index]], "")
        _clear_risk_table(document.tables[table_offset + remaining_index])


def _ensure_capacity_for_findings(
    document: DocumentObject,
    *,
    findings: list[Finding],
    paragraph_offset: int,
    table_offset: int,
    base_risk_heading_indices: list[int],
) -> list[int]:
    risk_heading_indices = list(base_risk_heading_indices)
    if len(findings) <= len(risk_heading_indices):
        return risk_heading_indices

    # Current template exposes 10 risk slots. Extra findings are not truncated:
    # clone the last placeholder heading + table to preserve the exact visual style.
    last_heading = document.paragraphs[paragraph_offset + risk_heading_indices[-1]]._p
    last_table = document.tables[table_offset + len(risk_heading_indices) - 1]._tbl
    body = document.element.body
    sect_pr = body.sectPr

    for _ in range(len(findings) - len(risk_heading_indices)):
        body.insert(body.index(sect_pr), deepcopy(last_heading))
        body.insert(body.index(sect_pr), deepcopy(last_table))
        risk_heading_indices.append(risk_heading_indices[-1] + 1)

    return risk_heading_indices


def _build_no_findings_placeholder(task: ScanTask, summary: dict[str, str]) -> Finding:
    return Finding(
        finding_id="finding-000",
        title="未发现风险项",
        severity="info",
        summary=summary.get("executive_summary", "本次扫描未发现可确认风险。"),
        evidence=summary.get("technical_analysis", f"扫描目标：{task.target}"),
        remediation=summary.get("recommendations", "继续保持现有安全控制，并定期复测。"),
    )


def _fill_risk_table(
    table: Table,
    *,
    task: ScanTask,
    finding: Finding,
    summary: dict[str, str],
) -> None:
    scan_time = task.created_at.astimezone().strftime("%Y-%m-%d %H:%M:%S")
    risk_level = SEVERITY_LABELS.get(finding.severity, finding.severity)

    _replace_cell_text(table.cell(0, 1), finding.title)
    _replace_cell_text(table.cell(1, 1), risk_level)
    _replace_cell_text(table.cell(2, 1), finding.summary)
    _replace_cell_text(
        table.cell(3, 1),
        "\n".join(
            [
                f"测试链接：{task.target}",
                "测试参数：-",
                f"功能点：{task.name}",
                f"测试时间：{scan_time}",
                f"测试过程：{summary.get('technical_analysis') or finding.summary}",
                f"漏洞证明：{finding.evidence}",
            ]
        ),
    )
    _replace_cell_text(
        table.cell(4, 1),
        "\n".join(
            [
                finding.summary,
                f"证据补充：{finding.evidence}",
            ]
        ),
    )
    _replace_cell_text(table.cell(5, 1), finding.remediation)


def _clear_risk_table(table: Table) -> None:
    for row_index in range(len(table.rows)):
        _replace_cell_text(table.cell(row_index, 1), "")


def _replace_cell_text(cell: _Cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return

    for paragraph_index, paragraph in enumerate(cell.paragraphs):
        _replace_paragraph_text(paragraph, text if paragraph_index == 0 else "")


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
        return

    paragraph.add_run(text)
