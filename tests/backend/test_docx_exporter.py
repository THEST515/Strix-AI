import sys
import unittest
from datetime import datetime, timezone
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.backend.domain.models import Finding, ScanReport, ScanTask, TaskStatus


class DocxExporterTests(unittest.TestCase):
    def test_builds_single_site_docx_from_template_without_losing_table_structure(self) -> None:
        from src.backend.services.docx_exporter import build_docx_export_bytes

        template_path = ROOT / "tests" / "fixtures" / "report_template.docx"
        task = ScanTask(
            task_id="task-001",
            name="授权测试站点",
            target="https://authorized-lab.example",
            scan_mode="deep",
            status=TaskStatus.COMPLETED,
            created_at=datetime(2026, 7, 7, tzinfo=timezone.utc),
        )
        report = ScanReport(
            task_id="task-001",
            findings=[
                Finding(
                    finding_id="finding-001",
                    title="个人资料接口越权访问",
                    severity="high",
                    summary="未授权用户可读取其他用户资料。",
                    evidence="替换用户编号后仍可返回他人资料字段。",
                    remediation="对对象读取增加服务端归属校验。",
                )
            ],
        )
        summary = {
            "executive_summary": "存在 1 项高危风险。",
            "technical_analysis": "接口对象访问控制缺失。",
            "recommendations": "优先修复接口鉴权。",
        }

        content = build_docx_export_bytes(
            template_path=template_path,
            task_exports=[{"task": task, "report": report, "summary": summary}],
        )

        with TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "single.docx"
            output_path.write_bytes(content)
            doc = Document(output_path)

        self.assertEqual(len(doc.tables), 10)
        self.assertIn("授权测试站点", [p.text for p in doc.paragraphs])
        self.assertIn("[高危]个人资料接口越权访问", [p.text for p in doc.paragraphs])
        self.assertEqual(doc.tables[0].cell(0, 1).text, "个人资料接口越权访问")
        self.assertEqual(doc.tables[0].cell(1, 1).text, "高危（已确认）")
        self.assertIn("未授权用户可读取其他用户资料。", doc.tables[0].cell(2, 1).text)
        self.assertIn("https://authorized-lab.example", doc.tables[0].cell(3, 1).text)
        self.assertIn("对对象读取增加服务端归属校验。", doc.tables[0].cell(5, 1).text)

    def test_builds_merged_docx_by_repeating_template_for_multiple_sites(self) -> None:
        from src.backend.services.docx_exporter import build_docx_export_bytes

        template_path = ROOT / "tests" / "fixtures" / "report_template.docx"
        task_a = ScanTask(
            task_id="task-001",
            name="站点 A",
            target="https://a.example",
            scan_mode="quick",
            status=TaskStatus.COMPLETED,
            created_at=datetime(2026, 7, 7, tzinfo=timezone.utc),
        )
        task_b = ScanTask(
            task_id="task-002",
            name="站点 B",
            target="https://b.example",
            scan_mode="quick",
            status=TaskStatus.COMPLETED,
            created_at=datetime(2026, 7, 7, tzinfo=timezone.utc),
        )
        report_a = ScanReport(
            task_id="task-001",
            findings=[
                Finding("finding-001", "风险 A", "medium", "摘要 A", "证据 A", "建议 A")
            ],
        )
        report_b = ScanReport(
            task_id="task-002",
            findings=[
                Finding("finding-002", "风险 B", "low", "摘要 B", "证据 B", "建议 B")
            ],
        )

        content = build_docx_export_bytes(
            template_path=template_path,
            task_exports=[
                {"task": task_a, "report": report_a, "summary": {}},
                {"task": task_b, "report": report_b, "summary": {}},
            ],
        )

        with TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "merged.docx"
            output_path.write_bytes(content)
            doc = Document(output_path)

        paragraph_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertIn("站点 A", paragraph_texts)
        self.assertIn("站点 B", paragraph_texts)
        self.assertIn("[中危]风险 A", paragraph_texts)
        self.assertIn("[低危]风险 B", paragraph_texts)
        self.assertGreaterEqual(len(doc.tables), 20)


if __name__ == "__main__":
    unittest.main()
